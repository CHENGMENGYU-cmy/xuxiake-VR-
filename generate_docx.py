"""
将徐霞客社区研究计划与研究方案 V2.0 转换为格式化的 Word 文档
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# ── 默认字体 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── 标题样式 ──
for i in range(1, 5):
    heading = doc.styles[f'Heading {i}']
    heading.font.name = 'Microsoft YaHei'
    heading.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    heading.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    if i == 1:
        heading.font.size = Pt(22)
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)
    elif i == 2:
        heading.font.size = Pt(16)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(8)
    elif i == 3:
        heading.font.size = Pt(13)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

# ── 工具函数 ──
def add_styled_table(headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后空行
    return table

def add_code_block(text):
    """添加代码块样式段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_body(text):
    """添加正文段落"""
    p = doc.add_paragraph(text)
    return p

def add_bullet(text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_separator():
    """添加分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 50)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('徐霞客社区')
run.font.size = Pt(38)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('研究计划与研究方案')
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x2d, 0x50, 0x8c)
run.font.name = 'Microsoft YaHei'
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in ['版本：V2.0', '日期：2026年7月20日', '状态：框架搭建完成，进入数据接入与测试阶段']:
    run = meta.add_run(line + '\n')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()

# ═══════════════════════════════════════════
# 导言：四个核心概念
# ═══════════════════════════════════════════
doc.add_heading('导言：四个核心概念', level=1)

add_body('在展开研究计划之前，先厘清本报告中四个贯穿始终的核心概念——它们回答不同的问题，不可混用。')

add_styled_table(
    ['概念', '回答的问题', '核心产出', '时间属性'],
    [
        ['任务（Task）', '要研究什么？达成什么目标？', '研究目标、关键成果、验收标准', '目标导向，不定时间'],
        ['方案（Solution）', '怎么实现？用什么技术路线？', '技术架构、接口设计、算法选型', '方向性，不定时间'],
        ['计划（Plan）', '什么时候做？谁来做？', '阶段划分、里程碑、资源分配', '严格时间线'],
        ['进度（Progress）', '做到哪了？还差什么？', '完成状态、遗留问题、下一步', '实时快照'],
    ]
)

add_body('四者关系：任务定义"做什么"，方案描述"怎么做"，计划安排"何时做"，进度呈现"做到哪"。任务驱动方案，方案拆分计划，计划追踪进度，进度反馈修正任务——这是一个持续循环的管理闭环。')

# ═══════════════════════════════════════════
# 一、项目概述
# ═══════════════════════════════════════════
doc.add_heading('一、项目概述', level=1)

doc.add_heading('1.1 项目定位', level=2)
add_body('徐霞客系统是一个以 VR/AR 内容为核心的沉浸式旅行社交分享平台，让用户以第一人称视角记录和分享旅行体验。平台融合了 Instagram 的社交分享、YouTube 的视频内容、小红书的笔记功能、Strava 的运动轨迹等平台特色，打造差异化的 VR 旅行社区。')

doc.add_heading('1.2 核心价值主张', level=2)
add_bullet('沉浸式体验：支持 VR360/VR180/空间视频，让用户足不出户体验世界各地风景')
add_bullet('内容驱动社交：围绕旅行内容建立社交关系，而非传统的聊天驱动')
add_bullet('AR 增强现实：通过 AR 眼镜数据接入，提供真实的增强现实旅行体验')
add_bullet('社区化运营：社群从群聊升级为内容社区，形成旅行爱好者聚集地')

# ═══════════════════════════════════════════
# 二、当前进度
# ═══════════════════════════════════════════
doc.add_heading('二、当前进度——框架搭建完成', level=1)

doc.add_heading('2.1 技术架构已就绪', level=2)
add_styled_table(
    ['层级', '技术选型', '状态'],
    [
        ['前端框架', 'Next.js 16 + React 19 + TypeScript', '运行中'],
        ['UI 组件库', 'shadcn/ui (Radix) + Tailwind CSS v4', '运行中'],
        ['状态管理', 'Zustand 5', '运行中'],
        ['后端框架', 'NestJS 11 + TypeScript', '运行中'],
        ['数据库', 'MySQL 8 + TypeORM', '运行中'],
        ['实时通信', 'Socket.IO (聊天) + SSE (通知)', '运行中'],
        ['认证', 'JWT 双 Token (Access 15min / Refresh 7d)', '运行中'],
    ]
)

doc.add_heading('2.2 已完成功能模块（八大系统）', level=2)
add_styled_table(
    ['序号', '模块', '完成度', '关键能力'],
    [
        ['1', '用户系统', '100%', '邮箱/手机注册登录、JWT双Token、个人主页、XXK编号、VR设备关联'],
        ['2', '内容发布', '100%', '6种帖子类型、VR格式、多媒体上传、位置标记、话题标签、路线创建器、攻略创建器、合集管理、草稿管理'],
        ['3', '信息流', '100%', '关注动态、探索发现、热门/最新/精选排序、游标分页、内容类型过滤、媒体发现'],
        ['4', '社交互动', '100%', '点赞、评论(嵌套回复)、关注/取关、分享、帖子编辑删除'],
        ['5', '实时通信', '100%', 'WebSocket私聊/群聊、媒体消息、消息反应、输入状态、已读回执、位置共享、广播消息、阅后即焚、AI助手、语音消息、VR预览'],
        ['6', '通知中心', '100%', 'SSE实时推送、通知分类、未读计数、标记已读'],
        ['7', '社群功能', '85%', '创建/加入/管理、动态流、公告系统、角色系统、挑战系统(基础)、兴趣标签'],
        ['8', '推荐系统', '70%', '多信号加权推荐(标签匹配/社交关系/地理匹配/活跃度)、用户推荐、社区推荐'],
    ]
)

doc.add_heading('2.3 数据库设计已完成', level=2)
add_body('28张数据表已定义并迁移至MySQL，覆盖用户、内容、社交、消息、社群五大域：')
add_bullet('用户域：users, user_follows, user_interests, interest_tags')
add_bullet('内容域：posts, media_items, comments, likes, topics, route_details, journeys, journey_stops, guide_details, collections, collection_posts')
add_bullet('社交域：notifications')
add_bullet('消息域：conversations, conversation_participants, messages, message_reactions, location_shares')
add_bullet('社群域：communities, community_tags, community_announcements, community_roles, community_challenges, community_challenge_entries')

doc.add_heading('2.4 开发进度判定', level=2)
add_body('当前阶段：框架搭建阶段已结束。已完成前后端技术栈搭建并稳定运行、28张数据表全部设计完成并迁移、8大功能系统开发基本完成、测试种子数据已准备（10个测试用户 + 6篇示例帖子 + 社区体验数据）。')
add_body('接下来需要进入的阶段：AR 眼镜真实数据接入 → 系统测试 → 公开上线。')

# ═══════════════════════════════════════════
# 三、研究任务
# ═══════════════════════════════════════════
doc.add_heading('三、研究任务——要研究什么', level=1)

doc.add_heading('3.1 任务体系总览', level=2)
add_body('五大核心研究任务构成递进式体系：')
add_bullet('任务1：AR 眼镜数据接入 — 实现 AR 眼镜与平台的数据对接，支持实时 AR 内容采集与上传')
add_bullet('任务2：AR 内容渲染与交互 — 在 Web 端实现高质量的 AR 内容渲染与沉浸式交互体验')
add_bullet('任务3：社区功能深化 — 将社群从基础框架升级为完整的旅行内容社区')
add_bullet('任务4：推荐算法优化 — 提升内容推荐的精准度和用户体验')
add_bullet('任务5：性能优化与质量保障 — 确保平台在高并发场景下的稳定性、流畅性和安全性')

doc.add_heading('3.2 任务详情与验收标准', level=2)

# 任务1
doc.add_heading('任务1：AR 眼镜数据接入', level=3)
add_body('研究问题：如何将 AR 眼镜产生的实时视频流、传感器数据、地理位置信息高效传输到平台后端？')
add_bullet('AR 眼镜数据格式调研（不同品牌/型号的数据输出差异）')
add_bullet('实时数据传输协议选型（WebSocket vs WebRTC vs 混合方案）')
add_bullet('数据预处理与压缩算法（视频编码、传感器数据降采样）')
add_bullet('边缘计算与云端协同架构（端侧预处理 vs 云端处理的分工）')

add_styled_table(
    ['验收项', '标准'],
    [
        ['支持 AR 眼镜品牌', '至少 3 种主流品牌'],
        ['视频流传输延迟', '< 200ms'],
        ['传感器数据同步误差', '< 50ms'],
        ['丢包率', '< 1%'],
        ['弱网环境', '可降级运行'],
    ]
)

# 任务2
doc.add_heading('任务2：AR 内容渲染与交互', level=3)
add_body('研究问题：如何在 Web 浏览器中高质量地渲染 AR 眼镜采集的内容，并提供沉浸式交互？')
add_bullet('WebXR API 集成方案研究（immersive-vr / immersive-ar 模式）')
add_bullet('3D 场景渲染引擎选型与优化（Three.js / CesiumJS / A-Frame）')
add_bullet('AR 标记点系统设计（地理信息叠加、距离感知显示/隐藏）')
add_bullet('多平台兼容性方案（PC 浏览器、移动端、VR 头显）')

add_styled_table(
    ['验收项', '标准'],
    [
        ['支持设备', 'Meta Quest / Pico / HTC Vive'],
        ['渲染帧率', '> 60fps'],
        ['3D 场景加载', '< 3s'],
        ['降级方案', '支持非 XR 模式的 2D 浏览'],
    ]
)

# 任务3
doc.add_heading('任务3：社区功能深化', level=3)
add_body('研究问题：如何在现有社群基础框架上，构建完整的旅行内容社区生态？')
add_bullet('社群挑战系统完善（排行榜算法、成就徽章体系、奖励机制）')
add_bullet('社群地图功能（旅行轨迹可视化、打卡点聚合、热门路线推荐）')
add_bullet('社群攻略/路线共享系统（协作编辑、版本管理）')
add_bullet('内容策展系统（精选内容、编辑推荐、话题聚合）')

add_styled_table(
    ['验收项', '标准'],
    [
        ['社群活跃度提升', '50%+'],
        ['挑战参与率', '> 30%'],
        ['内容策展覆盖', '> 20% 优质内容'],
    ]
)

# 任务4
doc.add_heading('任务4：推荐算法优化', level=3)
add_body('研究问题：如何基于多维行为信号，构建更精准的个性化推荐系统？')
add_bullet('行为信号采集体系设计（显式信号 + 隐式信号 + 社交信号）')
add_bullet('混合推荐算法研究（内容推荐 + 协同过滤 + 知识推荐 + 上下文推荐）')
add_bullet('实时推荐更新机制（流式处理、增量更新）')
add_bullet('负反馈学习与个性化调整')

add_styled_table(
    ['验收项', '标准'],
    [
        ['推荐点击率提升', '30%+'],
        ['用户平均浏览时长提升', '25%+'],
        ['推荐多样性指数', '> 0.7'],
    ]
)

# 任务5
doc.add_heading('任务5：性能优化与质量保障', level=3)
add_body('研究问题：如何确保平台在真实用户场景下的稳定性、流畅性和安全性？')
add_bullet('前端性能优化（懒加载、虚拟滚动、代码分割、PWA）')
add_bullet('后端性能优化（数据库索引优化、查询优化、Redis 缓存层）')
add_bullet('实时通信优化（WebSocket 连接池、消息压缩、优雅降级）')
add_bullet('安全审计（XSS/CSRF 防护、数据加密、权限控制审查）')

add_styled_table(
    ['验收项', '标准'],
    [
        ['首屏加载', '< 2s (P95)'],
        ['API 响应', '< 100ms (P95)'],
        ['WebSocket 消息延迟', '< 50ms'],
        ['安全', '无高危漏洞'],
    ]
)

# ═══════════════════════════════════════════
# 四、研究方案
# ═══════════════════════════════════════════
doc.add_heading('四、研究方案——怎么实现', level=1)

doc.add_heading('4.1 任务-方案映射', level=2)
add_styled_table(
    ['任务', '核心技术路线', '关键技术选型'],
    [
        ['AR 眼镜数据接入', 'WebSocket + WebRTC 混合传输', 'MediaStream API, IndexedDB'],
        ['AR 内容渲染', 'WebXR + Three.js + CesiumJS', 'WebGL 2.0, GPGPU'],
        ['社区功能深化', '游戏化设计 + 空间数据可视化', 'Leaflet, 排行榜算法'],
        ['推荐算法优化', '混合推荐 + 实时流处理', 'Redis Streams, 协同过滤'],
        ['性能优化', '缓存分层 + 懒加载 + 连接池', 'Redis, CDN, PWA'],
    ]
)

doc.add_heading('4.2 方案1：AR 眼镜数据接入', level=2)

doc.add_heading('数据采集层——统一抽象接口', level=3)
add_code_block('interface ARGlassesData {')
add_code_block('  deviceId: string;')
add_code_block('  timestamp: number;')
add_code_block('  videoStream: MediaStream;  // 通过 MediaStream API 统一接入')
add_code_block('  sensorData: {')
add_code_block('    gyroscope: { x: number; y: number; z: number };')
add_code_block('    accelerometer: { x: number; y: number; z: number };')
add_code_block('    magnetometer: { x: number; y: number; z: number };')
add_code_block('  };')
add_code_block('  location: {')
add_code_block('    latitude: number; longitude: number;')
add_code_block('    altitude: number; accuracy: number;')
add_code_block('  };')
add_code_block('  metadata: {')
add_code_block('    deviceModel: string; resolution: { width: number; height: number };')
add_code_block('    frameRate: number; batteryLevel: number;')
add_code_block('  };')
add_code_block('}')

doc.add_heading('数据传输架构', level=3)
add_body('采用 WebRTC + WebSocket 混合方案：')
add_bullet('WebRTC：负责视频流传输——低延迟、自适应码率、NAT 穿透')
add_bullet('WebSocket：负责控制信令和传感器数据——可靠传输、低开销')
add_bullet('弱网降级：自动切换离线缓存模式（IndexedDB），恢复网络后批量同步')

doc.add_heading('数据处理管线', level=3)
add_body('原始数据 → 预处理(降噪/去重) → 压缩(H.265/gzip) → 传输 → 云端解码 → 存储(OSS/CDN) → 分发')

doc.add_heading('4.3 方案2：AR 内容渲染', level=2)

doc.add_heading('WebXR 集成架构', level=3)
add_code_block('class XRSessionManager {')
add_code_block('  async startSession(mode: "immersive-vr" | "immersive-ar") {')
add_code_block('    const session = await navigator.xr.requestSession(mode);')
add_code_block('    session.requestAnimationFrame(this.renderLoop);')
add_code_block('  }')
add_code_block('  renderLoop(time, frame) {')
add_code_block('    const pose = frame.getViewerPose(referenceSpace);')
add_code_block('    // 同步 AR 眼镜姿态 → 渲染 3D 场景')
add_code_block('    this.renderer.render(this.scene, this.camera);')
add_code_block('  }')
add_code_block('}')

doc.add_heading('3D 渲染策略', level=3)
add_styled_table(
    ['内容类型', '渲染引擎', '说明'],
    [
        ['VR360/VR180 视频', 'Three.js (球体贴图)', '等距柱状投影映射到球体内表面'],
        ['3D 路线轨迹', 'CesiumJS', '地球尺度下的 GPS 轨迹可视化'],
        ['AR 标记点', 'Three.js (CSS3DRenderer)', '基于地理坐标叠加文字/图片/3D模型'],
        ['沉浸式场景', 'A-Frame (声明式)', '快速原型和简单场景'],
    ]
)

doc.add_heading('4.4 方案3：社区功能深化', level=2)

doc.add_heading('社群挑战系统——游戏化引擎', level=3)
add_code_block('-- 挑战表扩展')
add_code_block('ALTER TABLE community_challenges')
add_code_block('  ADD COLUMN challenge_type ENUM("photo","video","checkin","distance","altitude"),')
add_code_block('  ADD COLUMN target_value DECIMAL(10,2),')
add_code_block('  ADD COLUMN reward_points INT DEFAULT 0,')
add_code_block('  ADD COLUMN difficulty ENUM("easy","medium","hard","extreme");')
add_code_block('')
add_code_block('-- 成就徽章表')
add_code_block('CREATE TABLE achievement_badges (')
add_code_block('  id INT PRIMARY KEY AUTO_INCREMENT,')
add_code_block('  name VARCHAR(100), icon_url VARCHAR(500),')
add_code_block('  category ENUM("exploration","creation","social","challenge"),')
add_code_block('  rarity ENUM("common","rare","epic","legendary"),')
add_code_block('  criteria JSON')
add_code_block(');')

add_body('排行榜算法：综合得分 = Σ(挑战完成度 × 难度系数) + Σ(徽章稀有度分) + 近7天活跃度 × 0.3')

doc.add_heading('4.5 方案4：推荐算法优化', level=2)

doc.add_heading('混合推荐算法架构', level=3)
add_body('推荐引擎采用经典的三层架构——召回 → 排序 → 重排：')
add_bullet('召回层（多路召回，各取 TopN）：内容召回 + 协同召回 + 知识召回 + 上下文召回 + 热门召回')
add_bullet('排序层（精排）：特征工程 → 模型预测 CTR → 排序')
add_bullet('重排层（策略调整）：多样性打散、去重过滤、负反馈学习')

doc.add_heading('信号体系', level=3)
add_styled_table(
    ['信号类别', '权重', '采集内容'],
    [
        ['显式信号（主动行为）', '高', '点赞、评论、分享、关注、收藏'],
        ['隐式信号（被动行为）', '低（需去噪）', '浏览时长、滚动深度、点击路径、搜索词'],
        ['社交信号（关系网络）', '中', '共同好友、互动频率、社群重叠度、信任度'],
    ]
)

doc.add_heading('4.6 方案5：性能优化', level=2)

doc.add_heading('前端优化策略', level=3)
add_styled_table(
    ['优化项', '方案', '预期收益'],
    [
        ['图片懒加载', 'Intersection Observer + 渐进式加载', '首屏减少 60% 图片请求'],
        ['虚拟滚动', 'react-window 替换长列表', '万条数据 DOM 节点 < 30'],
        ['代码分割', 'Next.js Dynamic Import + Suspense', '首屏 JS 减 40%'],
        ['PWA', 'Service Worker + Workbox 缓存', '二次访问 < 1s'],
        ['预加载', '<link rel="preload"> 关键资源', 'FCP 提升 20%'],
    ]
)

doc.add_heading('后端缓存策略', level=3)
add_code_block('// Redis 三级缓存')
add_code_block('L1: 热点数据 (TTL 60s)  — 首页动态、热门帖子')
add_code_block('L2: 温数据 (TTL 300s)   — 帖子详情、用户主页')
add_code_block('L3: 冷数据 (TTL 3600s)  — 静态配置、历史归档')

# ═══════════════════════════════════════════
# 五、实施计划
# ═══════════════════════════════════════════
doc.add_heading('五、实施计划——什么时候做', level=1)

doc.add_heading('5.1 总体时间线（2026.8 — 2027.5）', level=2)

add_styled_table(
    ['阶段', '时间', '周期', '核心目标'],
    [
        ['第一阶段：AR数据接入', '2026年8月-9月', '8周', '完成 AR 眼镜数据接入，进行基础功能测试'],
        ['第二阶段：AR内容渲染', '2026年10月-11月', '8周', '实现高质量 AR 内容渲染与沉浸式交互'],
        ['第三阶段：社区深化', '2026年12月-2027年1月', '8周', '完善社区功能，系统集成测试'],
        ['第四阶段：推荐优化', '2027年2月-3月', '8周', '优化推荐算法，内部测试'],
        ['第五阶段：公测上线', '2027年4月-5月', '8周', '公开测试，正式面向用户推出'],
    ]
)

doc.add_heading('5.2 第一阶段：AR 数据接入与基础测试（2026年8-9月）', level=2)
add_styled_table(
    ['周次', '工作内容', '交付物'],
    [
        ['W1-W2', 'AR 眼镜数据接口调研与抽象层开发', 'AR 数据接入 SDK'],
        ['W3-W4', 'WebSocket + WebRTC 传输系统开发', '实时传输系统'],
        ['W5', '数据预处理与存储管线', '数据处理管线'],
        ['W6-W7', '基础功能测试与问题修复', '测试报告'],
        ['W8', '性能基准测试', '性能基准报告'],
    ]
)

doc.add_heading('5.3 第二阶段：AR 内容渲染（2026年10-11月）', level=2)
add_styled_table(
    ['周次', '工作内容', '交付物'],
    [
        ['W1-W3', 'WebXR 集成开发（Three.js + CesiumJS）', 'WebXR 渲染引擎'],
        ['W4-W5', '3D 场景渲染优化（帧率、加载时间）', '渲染性能报告'],
        ['W6-W7', 'AR 标记点系统开发', 'AR 标记点系统'],
        ['W8', '多平台兼容性测试', '兼容性测试报告'],
    ]
)

doc.add_heading('5.4 第三阶段：社区深化与集成测试（2026年12月-2027年1月）', level=2)
add_styled_table(
    ['周次', '工作内容', '交付物'],
    [
        ['W1-W3', '社群挑战系统完善（排行榜+徽章）', '挑战系统 2.0'],
        ['W4-W5', '社群地图 + 内容策展系统', '地图+策展系统'],
        ['W6-W7', '全系统集成测试', '集成测试报告'],
        ['W8', '压力测试与性能调优', '性能优化报告'],
    ]
)

doc.add_heading('5.5 第四阶段：推荐优化与内测（2027年2-3月）', level=2)
add_styled_table(
    ['周次', '工作内容', '交付物'],
    [
        ['W1-W2', '行为信号采集系统开发', '信号采集系统'],
        ['W3-W5', '混合推荐算法实现', '推荐引擎'],
        ['W6-W7', '内部测试与反馈收集', '内测报告'],
        ['W8', '算法调优与 A/B 测试框架搭建', 'A/B 测试框架'],
    ]
)

doc.add_heading('5.6 第五阶段：公开测试与上线（2027年4-5月）', level=2)
add_styled_table(
    ['周次', '工作内容', '交付物'],
    [
        ['W1-W2', '公开测试准备（环境、数据、监控）', '公测环境'],
        ['W3-W6', '公开测试执行（限量用户邀请）', '公测报告'],
        ['W7', '问题修复与最终优化', '正式版本'],
        ['W8', '正式上线与监控', '上线版本 + 监控面板'],
    ]
)

doc.add_heading('5.7 资源需求', level=2)

doc.add_heading('人力资源', level=3)
add_styled_table(
    ['角色', '人数', '主要职责'],
    [
        ['项目经理', '1', '整体协调、进度管理、风险控制'],
        ['前端开发', '3', 'AR 渲染、UI 组件、性能优化'],
        ['后端开发', '2', 'API 开发、数据处理、系统架构'],
        ['算法工程师', '1', '推荐算法、数据分析'],
        ['测试工程师', '2', '单元测试、集成测试、性能测试'],
        ['UI/UX 设计师', '1', '交互设计、AR 体验设计'],
    ]
)

doc.add_heading('硬件资源', level=3)
add_styled_table(
    ['资源', '数量', '用途'],
    [
        ['AR 眼镜（多品牌）', '3-5 台', '数据采集与兼容性测试'],
        ['VR 头显', '3 台', 'WebXR 兼容性测试'],
        ['测试手机', '3 台', '移动端适配测试'],
        ['云服务器', '2 台', '开发测试环境'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 六、测试计划
# ═══════════════════════════════════════════
doc.add_heading('六、测试计划', level=1)

doc.add_heading('6.1 测试金字塔', level=2)
add_styled_table(
    ['层级', '占比', '类型', '说明'],
    [
        ['L4', '5%', 'UAT 用户验收测试', '真实用户场景验证'],
        ['L3', '15%', '系统测试', '端到端功能、性能、安全'],
        ['L2', '30%', '集成测试', '模块间协作、API 联调'],
        ['L1', '50%', '单元测试', '函数/组件/服务正确性'],
    ]
)

doc.add_heading('6.2 各阶段测试重点', level=2)
add_styled_table(
    ['开发阶段', '测试重点'],
    [
        ['AR 数据接入', '数据格式解析、传输稳定性、断线重连'],
        ['AR 内容渲染', '渲染性能、多设备兼容、降级方案'],
        ['社区功能深化', '挑战逻辑、排行榜正确性、权限控制'],
        ['推荐算法优化', '推荐准确率、A/B 实验、用户体验'],
        ['公开测试', '真实场景、边界条件、大规模并发'],
    ]
)

doc.add_heading('6.3 性能指标', level=2)
add_styled_table(
    ['指标', '基准值', '目标值', '极限值'],
    [
        ['AR 传输延迟', '< 500ms', '< 200ms', '< 100ms'],
        ['AR 渲染帧率', '> 30fps', '> 60fps', '> 90fps'],
        ['首屏加载', '< 5s', '< 2s', '< 1s'],
        ['API 响应 (P95)', '< 500ms', '< 100ms', '< 50ms'],
        ['WebSocket 延迟', '< 200ms', '< 50ms', '< 20ms'],
        ['并发用户', '500', '2000', '5000'],
        ['系统可用性', '99%', '99.5%', '99.9%'],
    ]
)

# ═══════════════════════════════════════════
# 七、风险分析
# ═══════════════════════════════════════════
doc.add_heading('七、风险分析与应对', level=1)

doc.add_heading('7.1 技术风险', level=2)
add_styled_table(
    ['风险', '概率', '影响', '应对措施'],
    [
        ['AR 眼镜兼容性差异大', '高', '高', '统一抽象层；优先适配主流品牌；建立兼容性测试矩阵'],
        ['WebXR 浏览器支持不足', '中', '中', 'WebXR Polyfill 兜底；非 XR 降级模式；跟踪标准演进'],
        ['高并发 WebSocket 瓶颈', '中', '高', '连接池 + 消息压缩 + 优雅降级至轮询'],
        ['推荐算法冷启动', '高', '中', '内容+热门兜底；新用户兴趣问卷；渐进式学习'],
    ]
)

doc.add_heading('7.2 项目风险', level=2)
add_styled_table(
    ['风险', '概率', '影响', '应对措施'],
    [
        ['AR 接入复杂度超预期', '中', '高', '预留 20% 缓冲；分阶段交付；备选方案'],
        ['设备资源不足', '中', '中', '提前采购；设备共享；模拟数据先行'],
        ['用户 AR 接受度低', '低', '高', '渐进式引导；非 AR 替代体验；调研驱动迭代'],
    ]
)

doc.add_heading('7.3 市场风险', level=2)
add_styled_table(
    ['风险', '概率', '影响', '应对措施'],
    [
        ['竞品推出类似功能', '中', '中', '深耕旅行垂直领域；社区壁垒；持续创新'],
        ['AR 硬件普及度不足', '高', '中', '支持手机端创作；降低 AR 内容生产门槛'],
    ]
)

# ═══════════════════════════════════════════
# 八、成功标准
# ═══════════════════════════════════════════
doc.add_heading('八、成功标准与评估', level=1)

doc.add_heading('8.1 技术成功标准', level=2)
add_styled_table(
    ['指标', '目标值', '测量方法'],
    [
        ['AR 数据传输延迟', '< 200ms', '端到端延迟监控'],
        ['AR 渲染帧率', '> 60fps', 'requestAnimationFrame 统计'],
        ['首屏加载时间', '< 2s (P95)', 'Lighthouse / Web Vitals'],
        ['API 响应时间', '< 100ms (P95)', 'Prometheus + Grafana'],
        ['系统可用性', '> 99.5%', '健康检查 + 告警'],
        ['错误率', '< 1%', 'Sentry 异常监控'],
    ]
)

doc.add_heading('8.2 业务成功标准', level=2)
add_styled_table(
    ['指标', '上线 3 个月目标', '上线 6 个月目标'],
    [
        ['注册用户数', '5,000', '20,000'],
        ['日活跃用户 (DAU)', '500', '2,000'],
        ['内容发布总量', '3,000', '15,000'],
        ['AR 内容占比', '10%', '25%'],
        ['社群创建数', '50', '200'],
        ['用户留存率 (D30)', '30%', '45%'],
        ['用户满意度 NPS', '> 40', '> 50'],
    ]
)

doc.add_heading('8.3 持续评估机制', level=2)
add_bullet('周度：开发进度评审、技术风险扫描')
add_bullet('月度：性能基准测试、用户反馈分析')
add_bullet('里程碑：阶段交付物评审、Go/No-Go 决策')
add_bullet('上线后：7×24 监控、每日数据复盘、双周迭代')

# ═══════════════════════════════════════════
# 九、附录
# ═══════════════════════════════════════════
doc.add_heading('九、附录', level=1)

doc.add_heading('9.1 术语表', level=2)
add_styled_table(
    ['术语', '定义'],
    [
        ['AR', '增强现实 (Augmented Reality) — 在真实世界上叠加虚拟信息'],
        ['VR', '虚拟现实 (Virtual Reality) — 完全沉浸式的虚拟环境'],
        ['WebXR', 'W3C 标准的 Web 扩展现实 API，统一 VR/AR 的浏览器接口'],
        ['VR360', '360° 全景视频，等距柱状投影'],
        ['VR180', '180° 立体视频，前后双镜头拍摄'],
        ['WebRTC', 'Web Real-Time Communication，浏览器点对点实时通信'],
        ['SSE', 'Server-Sent Events，服务器向客户端单向推送事件'],
        ['PWA', 'Progressive Web App，渐进式 Web 应用'],
        ['P95/P99', '第 95/99 百分位数，衡量长尾延迟'],
    ]
)

doc.add_heading('9.2 参考文档', level=2)
add_bullet('WebXR Device API Specification: https://www.w3.org/TR/webxr/')
add_bullet('Three.js Documentation: https://threejs.org/docs/')
add_bullet('CesiumJS Documentation: https://cesium.com/cesiumjs/')
add_bullet('Socket.IO Documentation: https://socket.io/docs/v4/')
add_bullet('NestJS Documentation: https://docs.nestjs.com/')
add_bullet('Next.js Documentation: https://nextjs.org/docs')

# ── 保存 ──
output_path = r'D:\Other\Cluade CodeProjects\XuXiaKe\研究计划与研究方案.docx'
doc.save(output_path)
print(f'文档已保存到: {output_path}')
